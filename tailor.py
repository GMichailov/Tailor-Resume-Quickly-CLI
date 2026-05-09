import copy
import json
import os
from pathlib import Path
from urllib import error, request

from docx import Document
from dotenv import load_dotenv

from storage import (
    get_latest_const_data,
    get_resume_by_name,
    delete_bullet_point,
    insert_bullet_point,
    insert_skill_entry,
    insert_summary_entry,
    list_bullet_point_groups,
    list_bullet_points_for_group,
    list_project_entries,
    list_skill_entries,
    list_summary_entries,
    list_work_entries,
    update_bullet_point,
)


load_dotenv()
OPENROUTER_API_KEY = os.getenv("API_KEY")
MODEL="nvidia/nemotron-3-super-120b-a12b:free"


def get_job_description_input():
    job_description = input(
        """
        Paste the full job description below:
        """
    ).strip()

    if not job_description:
        raise ValueError("Job description cannot be empty.")

    return job_description

def call_openrouter_model(prompt, model=MODEL):
    if not OPENROUTER_API_KEY:
        raise ValueError("Missing API_KEY in .env")

    payload = {
        "model": model,
        "messages": [
            {"role": "user", "content": prompt},
        ],
    }

    req = request.Request(
        "https://openrouter.ai/api/v1/chat/completions",
        data=json.dumps(payload).encode("utf-8"),
        headers={
            "Authorization": f"Bearer {OPENROUTER_API_KEY}",
            "Content-Type": "application/json",
        },
        method="POST",
    )

    try:
        with request.urlopen(req) as response:
            response_body = response.read().decode("utf-8")
    except error.HTTPError as exc:
        error_body = exc.read().decode("utf-8", errors="replace")
        raise RuntimeError(
            f"OpenRouter request failed with status {exc.code}: {error_body}"
        ) from exc
    except error.URLError as exc:
        raise RuntimeError(f"Unable to reach OpenRouter: {exc.reason}") from exc

    response_json = json.loads(response_body)
    return response_json["choices"][0]["message"]["content"]

def extract_list_or_error(llm_response):
    response_text = llm_response.strip()

    opening_bracket_index = response_text.find("[")
    if opening_bracket_index == -1:
        raise ValueError("Unable to find the opening square bracket in the LLM response.")

    closing_bracket_index = response_text.find("]", opening_bracket_index)
    if closing_bracket_index == -1:
        raise ValueError("Unable to find the closing square bracket in the LLM response.")

    list_contents = response_text[opening_bracket_index + 1 : closing_bracket_index].strip()
    if not list_contents:
        raise ValueError("LLM response did not contain any usable keywords.")

    cleaned_keywords = []
    seen_keywords = set()

    for item in list_contents.split(", "):
        cleaned_item = item.strip().strip("\"'").strip()
        if not cleaned_item:
            continue
        if cleaned_item in seen_keywords:
            continue

        seen_keywords.add(cleaned_item)
        cleaned_keywords.append(cleaned_item)

    if not cleaned_keywords:
        raise ValueError("LLM response did not contain any usable keywords.")

    return cleaned_keywords


def allow_user_to_add_or_remove_keywords(keywords):
    updated_keywords = []
    seen_keywords = set()

    for keyword in keywords:
        cleaned_keyword = keyword.strip()
        if not cleaned_keyword or cleaned_keyword in seen_keywords:
            continue
        seen_keywords.add(cleaned_keyword)
        updated_keywords.append(cleaned_keyword)

    while True:
        print("\nCurrent ATS keywords:")
        for index, keyword in enumerate(updated_keywords, start=1):
            print(f"{index}. {keyword}")

        selection = input(
            """
            Keyword review
            1. Add keyword
            2. Remove keyword
            0. Finish
            >
            """
        ).strip()

        if selection == "0":
            return updated_keywords
        elif selection == "1":
            new_keyword = input(
                """
                Enter the keyword to add:
                """
            ).strip()

            if not new_keyword:
                print("Keyword cannot be empty.\n")
                continue
            if new_keyword in seen_keywords:
                print("That keyword is already in the list.\n")
                continue

            seen_keywords.add(new_keyword)
            updated_keywords.append(new_keyword)
            print("Keyword added.\n")
        elif selection == "2":
            if not updated_keywords:
                print("There are no keywords to remove.\n")
                continue

            removal_choice = input(
                """
                Enter the number of the keyword to remove:
                """
            ).strip()

            if not removal_choice.isdigit():
                print("That is not a valid selection.\n")
                continue

            removal_index = int(removal_choice)
            if removal_index < 1 or removal_index > len(updated_keywords):
                print("That is not a valid selection.\n")
                continue

            removed_keyword = updated_keywords.pop(removal_index - 1)
            seen_keywords.remove(removed_keyword)
            print(f'Removed "{removed_keyword}".\n')
        else:
            print("That is not a valid selection.\n")


def get_keywords_hit_for_text(ats_keywords, text):
    normalized_text = (text or "").lower()
    return {
        keyword for keyword in ats_keywords
        if keyword and keyword.lower() in normalized_text
    }


def get_skill_keyword_hits(ats_keywords):
    keyword_hits = {}

    for skills_id, description, skills in list_skill_entries():
        searchable_text = " ".join(part for part in (description, skills) if part)
        keyword_hits[skills_id] = get_keywords_hit_for_text(ats_keywords, searchable_text)

    return keyword_hits


def get_summary_keyword_hits(ats_keywords):
    keyword_hits = {}

    for summary_id, description, summary in list_summary_entries():
        searchable_text = " ".join(part for part in (description, summary) if part)
        keyword_hits[summary_id] = get_keywords_hit_for_text(ats_keywords, searchable_text)

    return keyword_hits


def get_project_cluster_keyword_hits(ats_keywords):
    keyword_hits = {}

    for project_id, title, start_date, end_date in list_project_entries():
        for cluster_id, parent_id, description in list_bullet_point_groups(project_id):
            bullet_points = list_bullet_points_for_group(cluster_id)
            bullet_text = " ".join(bullet_point[2] for bullet_point in bullet_points if bullet_point[2])
            searchable_text = " ".join(
                part for part in (title, description, bullet_text) if part
            )
            keyword_hits[(project_id, cluster_id)] = get_keywords_hit_for_text(
                ats_keywords,
                searchable_text,
            )

    return keyword_hits


def get_work_cluster_keyword_hits(ats_keywords):
    keyword_hits = {}

    for work_id, job_title, company, start_date, end_date in list_work_entries():
        for cluster_id, parent_id, description in list_bullet_point_groups(work_id):
            bullet_points = list_bullet_points_for_group(cluster_id)
            bullet_text = " ".join(bullet_point[2] for bullet_point in bullet_points if bullet_point[2])
            searchable_text = " ".join(
                part for part in (job_title, company, description, bullet_text) if part
            )
            keyword_hits[(work_id, cluster_id)] = get_keywords_hit_for_text(
                ats_keywords,
                searchable_text,
            )

    return keyword_hits


def select_best_resume_components(
    skill_keyword_hits,
    summary_keyword_hits,
    project_cluster_keyword_hits,
    work_cluster_keyword_hits,
):
    selected_skill = None
    if skill_keyword_hits:
        selected_skill = max(
            skill_keyword_hits.items(),
            key=lambda item: (len(item[1]), -item[0]),
        )

    selected_summary = None
    if summary_keyword_hits:
        selected_summary = max(
            summary_keyword_hits.items(),
            key=lambda item: (len(item[1]), -item[0]),
        )

    ranked_project_clusters = sorted(
        project_cluster_keyword_hits.items(),
        key=lambda item: (-len(item[1]), item[0][0], item[0][1]),
    )
    selected_project_clusters = []
    selected_project_ids = set()
    for (project_id, cluster_id), hits in ranked_project_clusters:
        if project_id in selected_project_ids:
            continue
        selected_project_ids.add(project_id)
        selected_project_clusters.append(((project_id, cluster_id), hits))
        if len(selected_project_clusters) == 2:
            break

    selected_work_clusters = {}
    work_clusters_by_work_id = {}
    for (work_id, cluster_id), hits in work_cluster_keyword_hits.items():
        work_clusters_by_work_id.setdefault(work_id, []).append(((work_id, cluster_id), hits))

    for work_id, clusters in work_clusters_by_work_id.items():
        selected_work_clusters[work_id] = sorted(
            clusters,
            key=lambda item: (-len(item[1]), item[0][1]),
        )[:2]

    return {
        "skill": selected_skill,
        "summary": selected_summary,
        "project_clusters": selected_project_clusters,
        "work_clusters": selected_work_clusters,
    }


def display_selected_resume_components(
    ats_keywords,
    skill_keyword_hits,
    summary_keyword_hits,
    project_cluster_keyword_hits,
    work_cluster_keyword_hits,
):
    selection = select_best_resume_components(
        skill_keyword_hits,
        summary_keyword_hits,
        project_cluster_keyword_hits,
        work_cluster_keyword_hits,
    )

    skill_lookup = {
        skills_id: (description, skills)
        for skills_id, description, skills in list_skill_entries()
    }
    summary_lookup = {
        summary_id: (description, summary)
        for summary_id, description, summary in list_summary_entries()
    }
    project_lookup = {
        project_id: (title, start_date, end_date)
        for project_id, title, start_date, end_date in list_project_entries()
    }
    work_lookup = {
        work_id: (job_title, company, start_date, end_date)
        for work_id, job_title, company, start_date, end_date in list_work_entries()
    }

    cluster_lookup = {}
    for project_id in project_lookup:
        for cluster_id, parent_id, description in list_bullet_point_groups(project_id):
            cluster_lookup[cluster_id] = description
    for work_id in work_lookup:
        for cluster_id, parent_id, description in list_bullet_point_groups(work_id):
            cluster_lookup[cluster_id] = description

    matched_keywords = set()

    print("\nRecommended Resume Combination:")

    print("\nTop Skill:")
    if selection["skill"] is None:
        print("No skill entries found.")
    else:
        skills_id, hits = selection["skill"]
        description, skills = skill_lookup.get(skills_id, (None, None))
        title = description or f"Skill entry {skills_id}"
        matched_keywords.update(hits)
        print(f"- {title}: {len(hits)} hits -> {sorted(hits)}")

    print("\nTop Summary:")
    if selection["summary"] is None:
        print("No summary entries found.")
    else:
        summary_id, hits = selection["summary"]
        description, summary = summary_lookup.get(summary_id, (None, None))
        title = description or f"Summary entry {summary_id}"
        matched_keywords.update(hits)
        print(f"- {title}: {len(hits)} hits -> {sorted(hits)}")

    print("\nTop Distinct Projects:")
    if not selection["project_clusters"]:
        print("No project clusters found.")
    for (project_id, cluster_id), hits in selection["project_clusters"]:
        title, start_date, end_date = project_lookup.get(project_id, (project_id, None, None))
        cluster_description = cluster_lookup.get(cluster_id) or cluster_id
        matched_keywords.update(hits)
        print(f"- {title} / {cluster_description}: {len(hits)} hits -> {sorted(hits)}")

    print("\nTop Experience Bullets Per Job:")
    if not selection["work_clusters"]:
        print("No work clusters found.")
    for work_id, work_row in work_lookup.items():
        job_title, company, start_date, end_date = work_row
        print(f"- {job_title} at {company}:")
        selected_clusters = selection["work_clusters"].get(work_id, [])
        if not selected_clusters:
            print("  No bullet point clusters found.")
            continue
        for (entry_work_id, cluster_id), hits in selected_clusters:
            cluster_description = cluster_lookup.get(cluster_id) or cluster_id
            matched_keywords.update(hits)
            print(f"  - {cluster_description}: {len(hits)} hits -> {sorted(hits)}")

    total_keywords = len(ats_keywords)
    missing_keywords = sorted(set(ats_keywords) - matched_keywords)

    print("\nATS Coverage Summary:")
    print(f"Matched {len(matched_keywords)} of {total_keywords} keywords.")
    print(f"Matched keywords: {sorted(matched_keywords)}")
    print(f"Missing keywords: {missing_keywords}")


def prompt_to_create_skill(ats_keywords, job_description):
    while True:
        selection = input(
            """
            Skill selection
            0. Go back
            1. Enter manually
            2. Generate with AI
            >
            """
        ).strip()

        if selection == "0":
            return
        if selection == "1":
            description = input(
                """
                Enter a description for this skill entry:
                """
            ).strip()
            skills = input(
                """
                Enter the skills text:
                """
            ).strip()
            if not skills:
                print("Skills text cannot be empty.\n")
                continue
            insert_skill_entry(description or None, skills)
            print("Skill entry saved.\n")
            return
        if selection == "2":
            prompt = (
                "Turn this list of ats_keywords from a job description in something useable in a resume directly."
                f"ATS keywords: {', '.join(ats_keywords)}.\n"
                f"Entire job description: {job_description}.\n"
                "Give your answer verbatim, say nothing other than the raw list of skills to maximize the ATS keywords being displayed in a non-suspicious way."
            )
            generated_skills = call_openrouter_model(prompt=prompt)
            description = input(
                """
                Enter a description for this generated skill entry:
                """
            ).strip()
            insert_skill_entry(description or None, generated_skills.strip())
            print("Generated skill entry saved.\n")
            return

        print("That is not a valid selection.\n")


def prompt_to_create_summary(ats_keywords, job_description):
    while True:
        selection = input(
            """
            Summary selection
            0. Go back
            1. Enter manually
            2. Generate with AI
            >
            """
        ).strip()

        if selection == "0":
            return
        if selection == "1":
            description = input(
                """
                Enter a description for this summary entry:
                """
            ).strip()
            summary = input(
                """
                Enter the summary text:
                """
            ).strip()
            if not summary:
                print("Summary text cannot be empty.\n")
                continue
            insert_summary_entry(description or None, summary)
            print("Summary entry saved.\n")
            return
        if selection == "2":
            prompt = (
                "Given a job description and its ats keywords, generate a non-suspicious summary for a resume while still trying to hit many keywords and facts about the job description."
                f"Job Description: {job_description}.\n"
                f"Keywords: {', '.join(ats_keywords)}.\n"
                "Your response should be the summary and nothing else. Your response verbatim will be copied and pasted."
            )
            generated_summary = call_openrouter_model(prompt=prompt)
            description = input(
                """
                Enter a description for this generated summary entry:
                """
            ).strip()
            insert_summary_entry(description or None, generated_summary.strip())
            print("Generated summary entry saved.\n")
            return

        print("That is not a valid selection.\n")


def prompt_user_to_select_skill(ats_keywords, job_description):
    while True:
        skill_entries = list_skill_entries()
        skill_keyword_hits = get_skill_keyword_hits(ats_keywords)

        print("\nSkill options:")
        for index, entry in enumerate(skill_entries, start=1):
            skills_id, description, skills = entry
            hits = sorted(skill_keyword_hits.get(skills_id, set()))
            title = description or f"Skill entry {skills_id}"
            print(f"{index}. {title}: {hits}")

        selection = input(
            """
            Enter the integer for the skill entry you want to select.
            Enter 0 to create a new skill entry.
            >
            """
        ).strip()

        if selection == "0":
            prompt_to_create_skill(ats_keywords, job_description)
            continue

        if not selection.isdigit():
            print("That is not a valid selection.\n")
            continue

        selected_index = int(selection)
        if selected_index < 1 or selected_index > len(skill_entries):
            print("That is not a valid selection.\n")
            continue

        selected_entry = skill_entries[selected_index - 1]
        print(f'Selected skill entry: {selected_entry[1] or f"Skill entry {selected_entry[0]}"}\n')
        return selected_entry[0]


def prompt_user_to_select_summary(ats_keywords, job_description):
    while True:
        summary_entries = list_summary_entries()
        summary_keyword_hits = get_summary_keyword_hits(ats_keywords)

        print("\nSummary options:")
        for index, entry in enumerate(summary_entries, start=1):
            summary_id, description, summary = entry
            hits = sorted(summary_keyword_hits.get(summary_id, set()))
            title = description or f"Summary entry {summary_id}"
            print(f"{index}. {title}: {hits}")

        selection = input(
            """
            Enter the integer for the summary entry you want to select.
            Enter 0 to create a new summary entry.
            >
            """
        ).strip()

        if selection == "0":
            prompt_to_create_summary(ats_keywords, job_description)
            continue

        if not selection.isdigit():
            print("That is not a valid selection.\n")
            continue

        selected_index = int(selection)
        if selected_index < 1 or selected_index > len(summary_entries):
            print("That is not a valid selection.\n")
            continue

        selected_entry = summary_entries[selected_index - 1]
        print(f'Selected summary entry: {selected_entry[1] or f"Summary entry {selected_entry[0]}"}\n')
        return selected_entry[0]


def prompt_user_to_select_project_clusters(ats_keywords, project_cluster_keyword_hits):
    selected_project_clusters = []

    print("\nProject cluster selection:")
    print("Enter 1 to include a cluster or 0 to skip it. You will be able to make edits in a moment.\n")

    for project_id, title, start_date, end_date in list_project_entries():
        print(f"Project: {title}")
        project_has_clusters = False

        for cluster_id, parent_id, description in list_bullet_point_groups(project_id):
            project_has_clusters = True
            hits = sorted(project_cluster_keyword_hits.get((project_id, cluster_id), set()))
            print(
                f"1. {description or cluster_id} - ({len(hits)} - {hits})"
            )

            while True:
                selection = input(
                    """
                    Enter 1 to include this cluster or 0 to skip it.
                    >
                    """
                ).strip()

                if selection == "0":
                    break
                if selection == "1":
                    selected_project_clusters.append((project_id, cluster_id))
                    print("Project cluster saved.\n")
                    break

                print("That is not a valid selection.\n")

        if not project_has_clusters:
            print("No bullet point clusters found.\n")

    return selected_project_clusters


def prompt_user_to_select_work_clusters(ats_keywords, work_cluster_keyword_hits):
    selected_work_clusters = []

    print("\nWork cluster selection:")
    print("Enter 1 to include a cluster or 0 to skip it. You will be able to make edits in a moment.\n")

    for work_id, job_title, company, start_date, end_date in list_work_entries():
        print(f"Work experience: {job_title}")
        work_has_clusters = False

        for cluster_id, parent_id, description in list_bullet_point_groups(work_id):
            work_has_clusters = True
            hits = sorted(work_cluster_keyword_hits.get((work_id, cluster_id), set()))
            print(
                f"1. {description or cluster_id} - ({len(hits)} - {hits})"
            )

            while True:
                selection = input(
                    """
                    Enter 1 to include this cluster or 0 to skip it.
                    >
                    """
                ).strip()

                if selection == "0":
                    break
                if selection == "1":
                    selected_work_clusters.append((work_id, cluster_id))
                    print("Work cluster saved.\n")
                    break

                print("That is not a valid selection.\n")

        if not work_has_clusters:
            print("No bullet point clusters found.\n")

    return selected_work_clusters


def calculate_current_selection_keyword_stats(
    ats_keywords,
    selected_summary,
    selected_skill,
    selected_project_clusters,
    selected_work_clusters,
):
    matched_keywords = set()

    if selected_summary is not None:
        matched_keywords.update(
            get_keywords_hit_for_text(ats_keywords, selected_summary.get("text", ""))
        )

    if selected_skill is not None:
        matched_keywords.update(
            get_keywords_hit_for_text(ats_keywords, selected_skill.get("text", ""))
        )

    for cluster in selected_project_clusters:
        cluster_text = " ".join(
            bullet_point["text"] for bullet_point in cluster.get("bullet_points", [])
        )
        cluster_hits = get_keywords_hit_for_text(
            ats_keywords,
            " ".join(
                part for part in (
                    cluster.get("title"),
                    cluster.get("cluster_description"),
                    cluster_text,
                ) if part
            ),
        )
        matched_keywords.update(cluster_hits)

    for cluster in selected_work_clusters:
        cluster_text = " ".join(
            bullet_point["text"] for bullet_point in cluster.get("bullet_points", [])
        )
        cluster_hits = get_keywords_hit_for_text(
            ats_keywords,
            " ".join(
                part for part in (
                    cluster.get("title"),
                    cluster.get("secondary_title"),
                    cluster.get("cluster_description"),
                    cluster_text,
                ) if part
            ),
        )
        matched_keywords.update(cluster_hits)

    return {
        "matched_keywords": matched_keywords,
        "missing_keywords": sorted(set(ats_keywords) - matched_keywords),
        "total_keywords": len(ats_keywords),
    }


def display_current_selection_and_missing_keywords(
    ats_keywords,
    selected_summary,
    selected_skill,
    selected_project_clusters,
    selected_work_clusters,
):
    stats = calculate_current_selection_keyword_stats(
        ats_keywords,
        selected_summary,
        selected_skill,
        selected_project_clusters,
        selected_work_clusters,
    )

    print("\nCurrent Selected Resume Content:")

    print("\nSummary:")
    if selected_summary is None:
        print("No summary selected.")
    else:
        title = selected_summary["description"] or f"Summary entry {selected_summary['id']}"
        hits = sorted(get_keywords_hit_for_text(ats_keywords, selected_summary.get("text", "")))
        print(f"- {title}: {hits}")

    print("\nWork Experiences:")
    if not selected_work_clusters:
        print("No work experience bullet groups selected.")
    else:
        work_clusters_by_work_id = {}
        for cluster in selected_work_clusters:
            work_clusters_by_work_id.setdefault(cluster["parent_id"], []).append(cluster)

        for work_id, clusters in work_clusters_by_work_id.items():
            print(f"- {clusters[0]['title']}")
            for cluster in clusters:
                cluster_text = " ".join(
                    bullet_point["text"] for bullet_point in cluster.get("bullet_points", [])
                )
                hits = sorted(
                    get_keywords_hit_for_text(
                        ats_keywords,
                        " ".join(
                            part for part in (
                                cluster.get("title"),
                                cluster.get("secondary_title"),
                                cluster.get("cluster_description"),
                                cluster_text,
                            ) if part
                        ),
                    )
                )
                print(f"  - {cluster['cluster_description'] or cluster['cluster_id']}: {hits}")

    print("\nProjects:")
    if not selected_project_clusters:
        print("No project bullet groups selected.")
    else:
        project_clusters_by_project_id = {}
        for cluster in selected_project_clusters:
            project_clusters_by_project_id.setdefault(cluster["parent_id"], []).append(cluster)

        for project_id, clusters in project_clusters_by_project_id.items():
            print(f"- {clusters[0]['title']}")
            for cluster in clusters:
                cluster_text = " ".join(
                    bullet_point["text"] for bullet_point in cluster.get("bullet_points", [])
                )
                hits = sorted(
                    get_keywords_hit_for_text(
                        ats_keywords,
                        " ".join(
                            part for part in (
                                cluster.get("title"),
                                cluster.get("cluster_description"),
                                cluster_text,
                            ) if part
                        ),
                    )
                )
                print(f"  - {cluster['cluster_description'] or cluster['cluster_id']}: {hits}")

    print("\nSelection ATS Coverage:")
    print(f"{len(stats['matched_keywords'])} / {stats['total_keywords']} keywords hit")
    print(f"Missing keywords: {stats['missing_keywords']}")


def save_modified_text_entry(insert_function, entry_type_label, text_value):
    save_choice = input(
        f"""
        Would you like to save this {entry_type_label} with a description? (y/n):
        """
    ).strip().lower()

    if save_choice != "y":
        return

    description = input(
        f"""
        Enter a description for this {entry_type_label}:
        """
    ).strip()
    insert_function(description or None, text_value)
    print(f"{entry_type_label.capitalize()} saved.\n")


def edit_selected_summary(ats_keywords, job_description, selected_summary):
    while True:
        selection = input(
            """
            Edit summary
            0. Back
            1. Manual
            2. AI
            >
            """
        ).strip()

        if selection == "0":
            return selected_summary
        if selection == "1":
            new_summary = input(
                """
                Enter the new summary text:
                """
            ).strip()
            if not new_summary:
                print("Summary text cannot be empty.\n")
                continue

            selected_summary["text"] = new_summary
            save_modified_text_entry(insert_summary_entry, "summary", new_summary)
            print("Summary updated.\n")
            continue
        if selection == "2":
            requested_changes = input(
                """
                What keywords would you like added or what changes should AI make?
                """
            ).strip()
            prompt = (
                "PROMPT\n"
                f"Job description: {job_description}\n"
                f"ATS keywords: {', '.join(ats_keywords)}\n"
                f"Current summary: {selected_summary['text']}\n"
                f"Requested changes: {requested_changes}\n"
            )
            generated_summary = call_openrouter_model(prompt=prompt).strip()
            selected_summary["text"] = generated_summary
            save_modified_text_entry(insert_summary_entry, "summary", generated_summary)
            print("Summary updated.\n")
            continue

        print("That is not a valid selection.\n")


def edit_selected_skill(ats_keywords, job_description, selected_skill):
    while True:
        selection = input(
            """
            Edit skills
            0. Back
            1. Manual
            2. AI
            >
            """
        ).strip()

        if selection == "0":
            return selected_skill
        if selection == "1":
            new_skills = input(
                """
                Enter the new skills text:
                """
            ).strip()
            if not new_skills:
                print("Skills text cannot be empty.\n")
                continue

            selected_skill["text"] = new_skills
            save_modified_text_entry(insert_skill_entry, "skills", new_skills)
            print("Skills updated.\n")
            continue
        if selection == "2":
            requested_changes = input(
                """
                What keywords would you like added or what changes should AI make?
                """
            ).strip()
            prompt = (
                "PROMPT\n"
                f"Job description: {job_description}\n"
                f"ATS keywords: {', '.join(ats_keywords)}\n"
                f"Current skills: {selected_skill['text']}\n"
                f"Requested changes: {requested_changes}\n"
            )
            generated_skills = call_openrouter_model(prompt=prompt).strip()
            selected_skill["text"] = generated_skills
            save_modified_text_entry(insert_skill_entry, "skills", generated_skills)
            print("Skills updated.\n")
            continue

        print("That is not a valid selection.\n")


def prompt_yes_no(message):
    return input(message).strip().lower() == "y"


def build_selected_cluster_objects(selected_clusters, entry_type):
    if entry_type == "project":
        parent_rows = {
            project_id: (title, start_date, end_date)
            for project_id, title, start_date, end_date in list_project_entries()
        }
    else:
        parent_rows = {
            work_id: (job_title, company, start_date, end_date)
            for work_id, job_title, company, start_date, end_date in list_work_entries()
        }

    cluster_objects = []
    for parent_id, cluster_id in selected_clusters:
        parent_details = parent_rows.get(parent_id)
        if parent_details is None:
            continue

        cluster_description = None
        for current_cluster_id, current_parent_id, description in list_bullet_point_groups(parent_id):
            if current_cluster_id == cluster_id:
                cluster_description = description
                break

        bullet_points = [
            {
                "id": bullet_id,
                "group_id": group_id,
                "text": description,
                "positional_order": positional_order,
            }
            for bullet_id, group_id, description, positional_order in list_bullet_points_for_group(cluster_id)
        ]

        cluster_objects.append(
            {
                "parent_id": parent_id,
                "cluster_id": cluster_id,
                "entry_type": entry_type,
                "title": parent_details[0],
                "secondary_title": parent_details[1] if entry_type == "work" else None,
                "start_date": parent_details[2] if entry_type == "work" else parent_details[1],
                "end_date": parent_details[3] if entry_type == "work" else parent_details[2],
                "cluster_description": cluster_description,
                "bullet_points": bullet_points,
            }
        )

    return cluster_objects


def edit_selected_cluster_bullets(selected_cluster):
    while True:
        print(f"\nEditing bullets for {selected_cluster['title']} / {selected_cluster['cluster_description'] or selected_cluster['cluster_id']}:")
        for index, bullet_point in enumerate(selected_cluster["bullet_points"], start=1):
            print(f"{index}. {bullet_point['text']}")

        selection = input(
            """
            0. Back
            1. Remove a bullet point
            2. Edit a bullet point
            3. Add a bullet point
            >
            """
        ).strip()

        if selection == "0":
            return selected_cluster
        if selection == "1":
            if not selected_cluster["bullet_points"]:
                print("No bullet points to remove.\n")
                continue

            bullet_choice = input(
                """
                Enter the integer of the bullet point to remove:
                """
            ).strip()
            if not bullet_choice.isdigit():
                print("That is not a valid selection.\n")
                continue

            bullet_index = int(bullet_choice)
            if bullet_index < 1 or bullet_index > len(selected_cluster["bullet_points"]):
                print("That is not a valid selection.\n")
                continue

            removed_bullet = selected_cluster["bullet_points"].pop(bullet_index - 1)
            if prompt_yes_no("Delete this bullet point from the DB as well? (y/n): "):
                delete_bullet_point(removed_bullet["id"])
            print("Bullet point removed.\n")
            continue
        if selection == "2":
            if not selected_cluster["bullet_points"]:
                print("No bullet points to edit.\n")
                continue

            bullet_choice = input(
                """
                Enter the integer of the bullet point to edit:
                """
            ).strip()
            if not bullet_choice.isdigit():
                print("That is not a valid selection.\n")
                continue

            bullet_index = int(bullet_choice)
            if bullet_index < 1 or bullet_index > len(selected_cluster["bullet_points"]):
                print("That is not a valid selection.\n")
                continue

            new_text = input(
                """
                Enter the new bullet point text:
                """
            ).strip()
            if not new_text:
                print("Bullet point text cannot be empty.\n")
                continue

            bullet_point = selected_cluster["bullet_points"][bullet_index - 1]
            bullet_point["text"] = new_text
            if prompt_yes_no("Overwrite this bullet point in the DB? (y/n): "):
                update_bullet_point(bullet_point["id"], new_text)
            print("Bullet point updated.\n")
            continue
        if selection == "3":
            new_text = input(
                """
                Enter the new bullet point text:
                """
            ).strip()
            if not new_text:
                print("Bullet point text cannot be empty.\n")
                continue

            next_position = len(selected_cluster["bullet_points"]) + 1
            new_bullet = {
                "id": None,
                "group_id": selected_cluster["cluster_id"],
                "text": new_text,
                "positional_order": next_position,
            }
            selected_cluster["bullet_points"].append(new_bullet)

            if prompt_yes_no("Save this bullet point for reuse in the DB? (y/n): "):
                insert_bullet_point(selected_cluster["cluster_id"], new_text, next_position)
                refreshed_bullets = list_bullet_points_for_group(selected_cluster["cluster_id"])
                selected_cluster["bullet_points"] = [
                    {
                        "id": bullet_id,
                        "group_id": group_id,
                        "text": description,
                        "positional_order": positional_order,
                    }
                    for bullet_id, group_id, description, positional_order in refreshed_bullets
                ]
            print("Bullet point added.\n")
            continue

        print("That is not a valid selection.\n")


def edit_selected_clusters(selected_clusters, entry_type_label):
    while True:
        print(f"\nSelected {entry_type_label}:")
        if not selected_clusters:
            print("No selected entries.\n")
        else:
            for index, cluster in enumerate(selected_clusters, start=1):
                print(f"{index}. {cluster['title']}")

        selection = input(
            f"""
            Enter the integer displayed before the {entry_type_label[:-1]}.
            0. Back
            >
            """
        ).strip()

        if selection == "0":
            return selected_clusters
        if not selection.isdigit():
            print("That is not a valid selection.\n")
            continue

        selected_index = int(selection)
        if selected_index < 1 or selected_index > len(selected_clusters):
            print("That is not a valid selection.\n")
            continue

        selected_cluster = selected_clusters[selected_index - 1]
        action = input(
            """
            0. Back
            1. Remove
            2. Edit
            >
            """
        ).strip()

        if action == "0":
            continue
        if action == "1":
            selected_clusters.pop(selected_index - 1)
            print("Selection removed.\n")
            continue
        if action == "2":
            selected_clusters[selected_index - 1] = edit_selected_cluster_bullets(selected_cluster)
            continue

        print("That is not a valid selection.\n")


def edit_selected_content(
    ats_keywords,
    job_description,
    selected_summary,
    selected_skill,
    selected_project_clusters,
    selected_work_clusters,
):
    while True:
        selection = input(
            """
            Edit selected content
            0. Exit
            1. Edit summary
            2. Edit projects
            3. Edit work
            4. Edit skills
            >
            """
        ).strip()

        if selection == "0":
            return (
                selected_summary,
                selected_skill,
                selected_project_clusters,
                selected_work_clusters,
            )
        if selection == "1":
            selected_summary = edit_selected_summary(
                ats_keywords,
                job_description,
                selected_summary,
            )
            continue
        if selection == "2":
            selected_project_clusters = edit_selected_clusters(
                selected_project_clusters,
                "projects",
            )
            continue
        if selection == "3":
            selected_work_clusters = edit_selected_clusters(
                selected_work_clusters,
                "work experiences",
            )
            continue
        if selection == "4":
            selected_skill = edit_selected_skill(
                ats_keywords,
                job_description,
                selected_skill,
            )
            continue

        print("That is not a valid selection.\n")


def review_and_edit_until_done(
    ats_keywords,
    job_description,
    selected_summary,
    selected_skill,
    selected_project_clusters,
    selected_work_clusters,
):
    while True:
        display_current_selection_and_missing_keywords(
            ats_keywords,
            selected_summary,
            selected_skill,
            selected_project_clusters,
            selected_work_clusters,
        )

        selection = input(
            """
            0. Finish and generate resume
            1. Continue editing
            >
            """
        ).strip()

        if selection == "0":
            return (
                selected_summary,
                selected_skill,
                selected_project_clusters,
                selected_work_clusters,
            )
        if selection == "1":
            (
                selected_summary,
                selected_skill,
                selected_project_clusters,
                selected_work_clusters,
            ) = edit_selected_content(
                ats_keywords,
                job_description,
                selected_summary,
                selected_skill,
                selected_project_clusters,
                selected_work_clusters,
            )
            continue

        print("That is not a valid selection.\n")


def replace_placeholder_in_paragraph(paragraph, placeholder, replacement_text):
    if placeholder not in paragraph.text:
        return False

    replacement_text = replacement_text or ""

    if not paragraph.runs:
        paragraph.text = paragraph.text.replace(placeholder, replacement_text)
        return True

    first_run = paragraph.runs[0]
    paragraph_text = paragraph.text.replace(placeholder, replacement_text)

    for run in paragraph.runs[1:]:
        run.text = ""
    first_run.text = paragraph_text
    return True


def replace_placeholder_everywhere(document, placeholder, replacement_text):
    replaced = False

    for paragraph in document.paragraphs:
        if replace_placeholder_in_paragraph(paragraph, placeholder, replacement_text):
            replaced = True

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if replace_placeholder_in_paragraph(paragraph, placeholder, replacement_text):
                        replaced = True

    return replaced


def insert_paragraph_after(paragraph, text=None, style=None):
    new_paragraph = paragraph.insert_paragraph_before(text or "", style)
    paragraph._element.addnext(new_paragraph._element)
    return new_paragraph


def find_paragraph_index_with_placeholder(document, placeholder):
    for index, paragraph in enumerate(document.paragraphs):
        if placeholder in paragraph.text:
            return index
    return None


def format_section_label(placeholder):
    token = placeholder.removeprefix("[[").removesuffix("]]").strip().lower()
    return " ".join(part.capitalize() for part in token.split("_"))


def build_date_text(start_date, end_date, use_present_for_missing_end=False):
    start_value = (start_date or "").strip()
    end_value = (end_date or "").strip()

    if not end_value and use_present_for_missing_end:
        end_value = "present"

    if start_value and end_value:
        return f"{start_value} - {end_value}"
    if start_value:
        return start_value
    if end_value:
        return end_value
    return ""


def apply_date_placeholder_paragraph(paragraph, start_date, end_date, use_present_for_missing_end=False):
    date_text = build_date_text(start_date, end_date, use_present_for_missing_end=use_present_for_missing_end)
    paragraph.text = date_text


def clone_repeating_block(document, anchor_placeholder, selected_clusters, use_present_for_missing_end=False):
    anchor_index = find_paragraph_index_with_placeholder(document, anchor_placeholder)
    if anchor_index is None:
        return False

    paragraphs = document.paragraphs
    anchor_paragraph = paragraphs[anchor_index]

    if anchor_index + 3 >= len(paragraphs):
        return False

    title_template_paragraph = paragraphs[anchor_index + 1]
    date_template_paragraph = paragraphs[anchor_index + 2]
    bullet_template_paragraph = paragraphs[anchor_index + 3]

    if not selected_clusters:
        anchor_paragraph.text = format_section_label(anchor_placeholder)
        title_template_paragraph.text = ""
        date_template_paragraph.text = ""
        bullet_template_paragraph.text = ""
        return True

    replace_placeholder_in_paragraph(
        anchor_paragraph,
        anchor_placeholder,
        format_section_label(anchor_placeholder),
    )

    current_anchor = date_template_paragraph
    first_entry = True

    for cluster in selected_clusters:
        if first_entry:
            block_title_paragraph = title_template_paragraph
            block_date_paragraph = date_template_paragraph
            first_bullet_paragraph = bullet_template_paragraph
            first_entry = False
        else:
            cloned_title = copy.deepcopy(title_template_paragraph._element)
            current_anchor._element.addnext(cloned_title)
            block_title_paragraph = document.paragraphs[document.paragraphs.index(current_anchor) + 1]

            cloned_date = copy.deepcopy(date_template_paragraph._element)
            block_title_paragraph._element.addnext(cloned_date)
            block_date_paragraph = document.paragraphs[document.paragraphs.index(block_title_paragraph) + 1]

            cloned_bullet = copy.deepcopy(bullet_template_paragraph._element)
            block_date_paragraph._element.addnext(cloned_bullet)
            first_bullet_paragraph = document.paragraphs[document.paragraphs.index(block_date_paragraph) + 1]

        replace_placeholder_in_paragraph(block_title_paragraph, "[[TITLE]]", cluster.get("title") or "")
        replace_placeholder_in_paragraph(
            block_title_paragraph,
            "[[COMPANY]]",
            cluster.get("secondary_title") or "",
        )

        apply_date_placeholder_paragraph(
            block_date_paragraph,
            cluster.get("start_date"),
            cluster.get("end_date"),
            use_present_for_missing_end=use_present_for_missing_end,
        )

        bullet_points = cluster.get("bullet_points", [])
        if bullet_points:
            replace_placeholder_in_paragraph(
                first_bullet_paragraph,
                "[[BULLET]]",
                bullet_points[0]["text"],
            )
            current_bullet_paragraph = first_bullet_paragraph
            for bullet_point in bullet_points[1:]:
                cloned_bullet = copy.deepcopy(bullet_template_paragraph._element)
                current_bullet_paragraph._element.addnext(cloned_bullet)
                current_bullet_paragraph = document.paragraphs[
                    document.paragraphs.index(current_bullet_paragraph) + 1
                ]
                replace_placeholder_in_paragraph(
                    current_bullet_paragraph,
                    "[[BULLET]]",
                    bullet_point["text"],
                )
        else:
            replace_placeholder_in_paragraph(first_bullet_paragraph, "[[BULLET]]", "")
            current_bullet_paragraph = first_bullet_paragraph

        current_anchor = current_bullet_paragraph

    return True


def generate_tailored_resume_docx(
    template_path,
    output_path,
    selected_summary,
    selected_skill,
    selected_project_clusters,
    selected_work_clusters,
):
    document = Document(template_path)
    latest_const_data = get_latest_const_data()

    scalar_data_map = {
        "[[NAME]]": latest_const_data[1] if latest_const_data else "",
        "[[EMAIL]]": latest_const_data[2] if latest_const_data else "",
        "[[LINKEDIN_URL]]": latest_const_data[3] if latest_const_data else "",
        "[[GITHUB_URL]]": latest_const_data[4] if latest_const_data else "",
        "[[PHONE_NUMBER]]": latest_const_data[5] if latest_const_data else "",
        "[[LOCATION]]": latest_const_data[6] if latest_const_data else "",
    }

    section_data_map = {
        "[[SUMMARY]]": selected_summary["text"] if selected_summary else "",
        "[[SKILLS]]": selected_skill["text"] if selected_skill else "",
        "[[CERTIFICATIONS]]": latest_const_data[7] if latest_const_data else "",
        "[[EDUCATION]]": latest_const_data[8] if latest_const_data else "",
        "[[EDUCATIONS]]": latest_const_data[8] if latest_const_data else "",
    }

    for placeholder, replacement_text in scalar_data_map.items():
        replace_placeholder_everywhere(document, placeholder, replacement_text or "")

    for placeholder, replacement_text in section_data_map.items():
        replace_placeholder_everywhere(
            document,
            placeholder,
            f"{format_section_label(placeholder)}\n{replacement_text or ''}".rstrip(),
        )

    clone_repeating_block(
        document,
        "[[EXPERIENCE]]",
        selected_work_clusters,
        use_present_for_missing_end=True,
    )
    clone_repeating_block(
        document,
        "[[PROJECTS]]",
        selected_project_clusters,
        use_present_for_missing_end=True,
    )

    document.save(output_path)


def prompt_for_resume_generation(
    selected_summary,
    selected_skill,
    selected_project_clusters,
    selected_work_clusters,
):
    resume_name = input(
        """
        Enter the name of the resume template you want to use:
        """
    ).strip()
    if not resume_name:
        print("Resume template name cannot be empty.\n")
        return

    resume_row = get_resume_by_name(resume_name)
    if not resume_row:
        print("Unable to find a resume template with that name.\n")
        return

    template_path = Path("files") / f"{resume_row[0]}.docx"
    if not template_path.exists():
        print("The resume template file could not be found.\n")
        return

    output_path_input = input(
        """
        Enter the absolute output path for the generated .docx file:
        """
    ).strip()
    output_path = Path(output_path_input)
    if not output_path.is_absolute():
        print("Please provide an absolute output path.\n")
        return
    if output_path.suffix.lower() != ".docx":
        print("Output path must end with .docx.\n")
        return

    generate_tailored_resume_docx(
        template_path=template_path,
        output_path=output_path,
        selected_summary=selected_summary,
        selected_skill=selected_skill,
        selected_project_clusters=selected_project_clusters,
        selected_work_clusters=selected_work_clusters,
    )
    print(f"Tailored resume saved to {output_path}\n")


def tailor_resume():
    # 1. Open up for for job description.
    job_description = get_job_description_input()
    # 2. Call ATS Parsing function with AI and get it returned as a list.
    llm_response = call_openrouter_model(
        prompt=(
            "You are a model that acts as an ATS Parser. You find all of the keywords or skills in a job description."
            "Given the following job description, identify these. These come from the resonsibilities, minimum qualifications, and preferred qualifications sections or their equivalents."
            f"{job_description}"
            "\nReturn your answer as a list wrapped in square brackets, separated by commas and spaces, and NO QUOTATION MARKS."
        )
    )
    ats_keywords = extract_list_or_error(llm_response)
    # 3. Allow user to see and include anything they believe is missing.
    ats_keywords = allow_user_to_add_or_remove_keywords(ats_keywords)
    # 4. Call function that goes through their skills, summaries, projects, and work experience and associates each with what tags each checks off.
    skill_keyword_hits = get_skill_keyword_hits(ats_keywords)
    summary_keyword_hits = get_summary_keyword_hits(ats_keywords)
    project_cluster_keyword_hits = get_project_cluster_keyword_hits(ats_keywords)
    work_cluster_keyword_hits = get_work_cluster_keyword_hits(ats_keywords)
    # 5. Calculate which combination of skill, summary, projects (allow user to pick a max), and work experiences (use all) checks off the highest number of qualifications. (Done in function below)
    display_selected_resume_components(
        ats_keywords,
        skill_keyword_hits,
        summary_keyword_hits,
        project_cluster_keyword_hits,
        work_cluster_keyword_hits,
    )
    # 6. Display the keyword scores of each and allow user to select (even if not recommended ones).
    selected_skill_id = prompt_user_to_select_skill(ats_keywords, job_description)
    selected_summary_id = prompt_user_to_select_summary(ats_keywords, job_description)

    selected_skill_row = next(
        (
            entry for entry in list_skill_entries()
            if entry[0] == selected_skill_id
        ),
        None,
    )
    selected_summary_row = next(
        (
            entry for entry in list_summary_entries()
            if entry[0] == selected_summary_id
        ),
        None,
    )

    selected_skill = None
    if selected_skill_row is not None:
        selected_skill = {
            "id": selected_skill_row[0],
            "description": selected_skill_row[1],
            "text": selected_skill_row[2] or "",
        }

    selected_summary = None
    if selected_summary_row is not None:
        selected_summary = {
            "id": selected_summary_row[0],
            "description": selected_summary_row[1],
            "text": selected_summary_row[2] or "",
        }
    selected_project_clusters = prompt_user_to_select_project_clusters(
        ats_keywords,
        project_cluster_keyword_hits,
    )
    selected_work_clusters = prompt_user_to_select_work_clusters(
        ats_keywords,
        work_cluster_keyword_hits,
    )

    selected_project_clusters = build_selected_cluster_objects(
        selected_project_clusters,
        "project",
    )
    selected_work_clusters = build_selected_cluster_objects(
        selected_work_clusters,
        "work",
    )
    # 7. Display to user which keywords are still missing.
    # 8. Give them the option to rewrite certain sections or bullet points with AI to include the missing keywords or to do it themselves.
    (
        selected_summary,
        selected_skill,
        selected_project_clusters,
        selected_work_clusters,
    ) = review_and_edit_until_done(
        ats_keywords,
        job_description,
        selected_summary,
        selected_skill,
        selected_project_clusters,
        selected_work_clusters,
    )
    # 10. Edit parent template in-place.
    # 11. Ask them where they would like this file saved and what it should be called.
    prompt_for_resume_generation(
        selected_summary,
        selected_skill,
        selected_project_clusters,
        selected_work_clusters,
    )
    return job_description
